"""
报告生成模块 - 基于 verified_findings artifact 生成 Word 报告
"""

import json
import os
import re
from datetime import datetime
from html import escape
from typing import Annotated, Dict, List, Optional

from core import namespace, tool, toolset

namespace()


@toolset()
class ReportGenerator:
    """生成包含漏洞详情和截图的 Word 渗透测试报告"""

    SEVERITY_ORDER = {"严重": 0, "高危": 1, "中危": 2, "低危": 3, "信息": 4}
    TARGET_PROFILE_LABELS = {
        "target": "目标 URL",
        "site_type": "站点类型",
        "application": "站点类型",
        "cms": "CMS",
        "theme": "主题/版本",
        "template": "主题/版本",
        "cdn": "CDN/WAF",
        "waf": "CDN/WAF",
        "edge_network": "CDN/WAF",
        "admin_username": "管理员用户名",
        "admin_user": "管理员用户名",
        "username": "管理员用户名",
    }

    def __init__(self, workspace: str = "/home/ubuntu/Workspace"):
        self.workspace = workspace
        self.screenshots_dir = os.path.join(workspace, "screenshots")
        os.makedirs(self.workspace, exist_ok=True)
        os.makedirs(self.screenshots_dir, exist_ok=True)

    def _safe_filename(self, value: str) -> str:
        sanitized = re.sub(r"[^A-Za-z0-9._-]+", "_", value).strip("._")
        return sanitized or "report"

    def _write_final_report_reference(self, report_path: str, report_type: str, target: str, total_findings: int) -> None:
        artifacts_dir = os.path.join(self.workspace, "intentlang", "artifacts")
        os.makedirs(artifacts_dir, exist_ok=True)
        artifact_path = os.path.join(artifacts_dir, "final_report_reference.json")
        now = datetime.now().isoformat(timespec="seconds")
        payload = {
            "artifact": "final_report_reference",
            "items": [
                {
                    "path": report_path,
                    "type": report_type,
                    "target": target,
                    "total_findings": total_findings,
                    "recorded_at": now,
                }
            ],
            "updated_at": now,
        }
        with open(artifact_path, "w", encoding="utf-8") as f:
            json.dump(payload, f, indent=2, ensure_ascii=False)
            f.write("\n")

    def _normalize_findings(self, findings: List[Dict]) -> List[Dict]:
        if not findings:
            findings = self._read_verified_findings_artifact()
        screenshot_indexes = self._candidate_screenshot_indexes()
        return self._sort_findings([self._coerce_finding(finding, screenshot_indexes) for finding in findings])

    def _report_filename(self, target: str, extension: str) -> str:
        safe_target = self._safe_filename(target)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        return f"Pentest_Report_{safe_target}_{timestamp}.{extension}"

    def _read_run_metadata(self) -> dict:
        metadata_path = os.path.join(self.workspace, "intentlang", "metadata", "run.json")
        if not os.path.exists(metadata_path):
            return {}
        with open(metadata_path, "r", encoding="utf-8") as f:
            return json.load(f)

    def _stringify_profile_value(self, value) -> str:
        if value is None:
            return ""
        if isinstance(value, list):
            return "、".join(str(item) for item in value if item is not None).strip()
        if isinstance(value, dict):
            return json.dumps(value, ensure_ascii=False)
        return str(value).strip()

    def _profile_label(self, key: str) -> str:
        if key in self.TARGET_PROFILE_LABELS:
            return self.TARGET_PROFILE_LABELS[key]
        return key.replace("_", " ").strip() or "附加信息"

    def _target_profile_rows(self, target: str) -> list[tuple[str, str]]:
        run_metadata = self._read_run_metadata()
        profile = {}
        for field_name in ("target_profile", "target_info"):
            field_value = run_metadata.get(field_name)
            if isinstance(field_value, dict):
                profile.update(field_value)

        for metadata_key in ("site_type", "application", "cms", "theme", "template", "cdn", "waf", "edge_network", "admin_username", "admin_user", "username"):
            if metadata_key in run_metadata and metadata_key not in profile:
                profile[metadata_key] = run_metadata[metadata_key]

        rows: list[tuple[str, str]] = [("目标 URL", target)]
        used_labels = {"目标 URL"}
        for key, value in profile.items():
            text = self._stringify_profile_value(value)
            if not text:
                continue
            label = self._profile_label(str(key))
            if label in used_labels:
                continue
            used_labels.add(label)
            rows.append((label, text))
        return rows

    def _read_verified_findings_artifact(self) -> List[Dict]:
        artifact_path = os.path.join(self.workspace, "intentlang", "artifacts", "verified_findings.json")
        if not os.path.exists(artifact_path):
            return []
        with open(artifact_path, "r", encoding="utf-8") as f:
            payload = json.load(f)
        return payload.get("items", [])

    def _read_candidate_evidence_artifact(self) -> List[Dict]:
        artifact_path = os.path.join(self.workspace, "intentlang", "artifacts", "candidate_evidence.json")
        if not os.path.exists(artifact_path):
            return []
        with open(artifact_path, "r", encoding="utf-8") as f:
            payload = json.load(f)
        return payload.get("items", [])

    def _normalize_lookup_key(self, value: str) -> str:
        return " ".join(str(value or "").strip().lower().split())

    def _candidate_screenshot_indexes(self) -> tuple[dict[str, str], dict[str, str], dict[str, str]]:
        by_evidence_id: dict[str, str] = {}
        by_finding_id: dict[str, str] = {}
        by_related_finding: dict[str, str] = {}
        for item in self._read_candidate_evidence_artifact():
            if item.get("kind") != "screenshot":
                continue
            path = str(item.get("path") or item.get("screenshot_path") or "").strip()
            evidence_id = self._normalize_lookup_key(item.get("evidence_id", ""))
            finding_id = self._normalize_lookup_key(item.get("finding_id", ""))
            related_finding = self._normalize_lookup_key(item.get("related_finding", ""))
            if not path or not os.path.exists(path):
                continue
            if evidence_id:
                by_evidence_id[evidence_id] = path
            if finding_id:
                by_finding_id[finding_id] = path
            if related_finding:
                by_related_finding[related_finding] = path
        return by_evidence_id, by_finding_id, by_related_finding

    def _finding_title(self, finding: Dict) -> str:
        return str(finding.get("name") or finding.get("title") or "未命名漏洞")

    def _resolve_screenshot_path(
        self,
        finding: Dict,
        screenshot_indexes: tuple[dict[str, str], dict[str, str], dict[str, str]] | None = None,
    ) -> str:
        if finding.get("screenshot_path"):
            return finding["screenshot_path"]
        if not screenshot_indexes:
            return ""

        by_evidence_id, by_finding_id, by_related_finding = screenshot_indexes
        evidence_id = self._normalize_lookup_key(finding.get("evidence_id", ""))
        if evidence_id and evidence_id in by_evidence_id:
            return by_evidence_id[evidence_id]

        finding_id = self._normalize_lookup_key(finding.get("finding_id", ""))
        if finding_id and finding_id in by_finding_id:
            return by_finding_id[finding_id]

        return by_related_finding.get(self._normalize_lookup_key(self._finding_title(finding)), "")

    def _coerce_finding(
        self,
        finding: Dict,
        screenshot_indexes: tuple[dict[str, str], dict[str, str], dict[str, str]] | None = None,
    ) -> Dict:
        title = self._finding_title(finding)
        resolved_screenshot = self._resolve_screenshot_path(finding, screenshot_indexes)
        evidence = finding.get("evidence")
        if not evidence:
            evidence = finding.get("test_process") or finding.get("evidence_summary") or finding.get("summary", "")
        impact = finding.get("impact")
        if not impact:
            impact = finding.get("risk_analysis") or finding.get("evidence_summary") or evidence or finding.get("summary", "")
        return {
            "name": title,
            "title": title,
            "severity": finding.get("severity", "信息"),
            "description": finding.get("description", finding.get("summary", "无描述")),
            "evidence": evidence,
            "impact": impact,
            "remediation": finding.get("remediation", "暂无修复建议"),
            "screenshot_path": resolved_screenshot,
            "type": finding.get("type", "其他"),
            "url": finding.get("url", finding.get("vuln_url", "")),
            "control_point": finding.get("control_point", ""),
            "evaluation_unit": finding.get("evaluation_unit", ""),
            "risk_analysis": finding.get("risk_analysis", impact),
            "test_process": finding.get("test_process", evidence),
            "vuln_code": finding.get("vuln_code", ""),
            "target": finding.get("target", ""),
            "finding_id": finding.get("finding_id", ""),
            "evidence_id": finding.get("evidence_id", ""),
        }

    def _count_severities(self, findings: List[Dict]) -> dict[str, int]:
        counts = {"严重": 0, "高危": 0, "中危": 0, "低危": 0, "信息": 0}
        for finding in findings:
            severity = finding.get("severity", "信息")
            if severity in counts:
                counts[severity] += 1
        return counts

    def _sort_findings(self, findings: List[Dict]) -> List[Dict]:
        return sorted(
            findings,
            key=lambda finding: (
                self.SEVERITY_ORDER.get(finding.get("severity", "信息"), 99),
                self._finding_title(finding),
            ),
        )

    def _set_cell_text(self, cell, text: str) -> None:
        cell.text = str(text)

    def _set_merged_row_content(self, table, row_idx: int, text: str) -> None:
        row = table.rows[row_idx]
        self._set_cell_text(row.cells[1], text)

    def _clear_cell(self, cell) -> None:
        cell.text = ""

    def _append_text_to_cell(self, cell, text: str) -> None:
        self._clear_cell(cell)
        lines = [line for line in str(text).splitlines()] or [""]
        for idx, line in enumerate(lines):
            paragraph = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
            paragraph.add_run(line)

    def _append_image_to_cell(self, cell, image_path: str, width_inches: float = 5.5) -> None:
        from docx.shared import Inches

        paragraph = cell.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(image_path, width=Inches(width_inches))

    def _report_times(self) -> tuple[str, str]:
        now = datetime.now()
        run_metadata = self._read_run_metadata()
        created_at = run_metadata.get("created_at", "")
        start_time = created_at.replace("T", " ").split("+")[0] if created_at else now.strftime("%Y-%m-%d %H:%M:%S")
        end_time = now.strftime("%Y-%m-%d %H:%M:%S")
        return start_time, end_time

    def _set_table_header(self, table, headers: list[str]) -> None:
        for index, value in enumerate(headers):
            self._set_cell_text(table.rows[0].cells[index], value)

    def _append_finding_table(self, doc, index: int, finding: Dict) -> None:
        title = self._finding_title(finding)
        control_point = finding.get("control_point") or "待补充"
        evaluation_unit = finding.get("evaluation_unit") or "待补充"
        severity = finding.get("severity", "信息")
        vuln_code = finding.get("vuln_code") or f"VUL-AUTO-{index:02d}"
        evidence = finding.get("test_process") or finding.get("evidence_summary") or "待补充"
        risk = finding.get("risk_analysis") or "待补充"
        description = finding.get("description") or finding.get("summary") or "待补充"
        remediation = finding.get("remediation") or "待补充"
        screenshot_path = finding.get("screenshot_path", "")
        vuln_url = finding.get("vuln_url", finding.get("url", finding.get("target", "")))

        doc.add_heading(f"{index}. {title}", level=2)
        table = doc.add_table(rows=8, cols=2)
        table.style = "Table Grid"
        self._set_table_header(table, ["字段", "内容"])

        fields = [
            ("漏洞编号", vuln_code),
            ("漏洞名称", title),
            ("安全控制点", control_point),
            ("测评单元", evaluation_unit),
            ("风险等级", severity),
            ("漏洞描述", description),
            ("漏洞链接", vuln_url),
        ]
        for row_index, (label, value) in enumerate(fields, start=1):
            self._set_cell_text(table.rows[row_index].cells[0], label)
            self._set_cell_text(table.rows[row_index].cells[1], value)

        detail_table = doc.add_table(rows=3, cols=2)
        detail_table.style = "Table Grid"
        self._set_table_header(detail_table, ["字段", "内容"])

        self._set_cell_text(detail_table.rows[1].cells[0], "测试过程")
        process_cell = detail_table.rows[1].cells[1]
        process_text = evidence or "已完成漏洞验证，详见截图与相关证据。"
        self._append_text_to_cell(process_cell, process_text)
        if screenshot_path and os.path.exists(screenshot_path):
            try:
                self._append_image_to_cell(process_cell, screenshot_path)
            except Exception as e:
                process_cell.add_paragraph(f"[截图插入失败: {e}]")
        elif screenshot_path:
            process_cell.add_paragraph(f"[截图文件不存在: {screenshot_path}]")

        self._set_cell_text(detail_table.rows[2].cells[0], "风险分析")
        self._set_cell_text(detail_table.rows[2].cells[1], risk)

        remediation_table = doc.add_table(rows=2, cols=2)
        remediation_table.style = "Table Grid"
        self._set_table_header(remediation_table, ["字段", "内容"])
        self._set_cell_text(remediation_table.rows[1].cells[0], "修复建议")
        self._set_cell_text(remediation_table.rows[1].cells[1], remediation)
        doc.add_paragraph()

    def _build_default_report(self, target: str, findings: List[Dict], report_title: str):
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        start_time, end_time = self._report_times()
        severity_counts = self._count_severities(findings)
        target_profile_rows = self._target_profile_rows(target)

        title = doc.add_heading(report_title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        time_para = doc.add_paragraph()
        time_para.add_run("测试时间：").bold = True
        time_para.add_run(f"{start_time} 至 {end_time}")

        summary = doc.add_paragraph()
        summary.add_run("执行摘要：").bold = True
        summary.add_run(f"共确认 {len(findings)} 个已验证安全发现。")

        doc.add_heading("目标信息", level=1)
        target_table = doc.add_table(rows=len(target_profile_rows) + 1, cols=2)
        target_table.style = "Table Grid"
        self._set_table_header(target_table, ["字段", "内容"])
        for row_index, (label, value) in enumerate(target_profile_rows, start=1):
            self._set_cell_text(target_table.rows[row_index].cells[0], label)
            self._set_cell_text(target_table.rows[row_index].cells[1], value)

        doc.add_heading("风险汇总", level=1)
        summary_table = doc.add_table(rows=6, cols=2)
        summary_table.style = "Table Grid"
        self._set_table_header(summary_table, ["风险等级", "数量"])
        summary_rows = [
            ("严重", str(severity_counts["严重"])),
            ("高危", str(severity_counts["高危"])),
            ("中危", str(severity_counts["中危"])),
            ("低危", str(severity_counts["低危"])),
            ("信息", str(severity_counts["信息"])),
        ]
        for row_index, (severity, count) in enumerate(summary_rows, start=1):
            self._set_cell_text(summary_table.rows[row_index].cells[0], severity)
            self._set_cell_text(summary_table.rows[row_index].cells[1], count)

        doc.add_heading("详细发现", level=1)
        for index, finding in enumerate(findings, start=1):
            self._append_finding_table(doc, index, finding)

        doc.add_heading("测试结论", level=1)
        doc.add_paragraph("本次评估已基于 verified_findings artifact 生成正式报告。建议按风险等级优先修复高危和中危问题，并复测验证。")
        return doc

    def _build_markdown_report(
        self,
        target: str,
        findings: List[Dict],
        report_title: str,
        fallback_reason: str = "",
    ) -> str:
        start_time, end_time = self._report_times()
        severity_counts = self._count_severities(findings)
        target_profile_rows = self._target_profile_rows(target)
        lines = [
            f"# {report_title}",
            "",
            f"- 测试时间: {start_time} 至 {end_time}",
            f"- 已验证发现数量: {len(findings)}",
            "",
            "## 目标信息",
            "",
            "| 字段 | 内容 |",
            "| --- | --- |",
        ]
        lines.extend(f"| {label} | {value} |" for label, value in target_profile_rows)
        if fallback_reason:
            lines.append(f"- 说明: {fallback_reason}")
        lines.extend(
            [
                "",
                "## 风险汇总",
                "",
                "| 风险等级 | 数量 |",
                "| --- | --- |",
                f"| 严重 | {severity_counts['严重']} |",
                f"| 高危 | {severity_counts['高危']} |",
                f"| 中危 | {severity_counts['中危']} |",
                f"| 低危 | {severity_counts['低危']} |",
                f"| 信息 | {severity_counts['信息']} |",
                "",
                "## 详细发现",
                "",
            ]
        )

        if not findings:
            lines.extend(["当前未发现已验证漏洞。", ""])
        else:
            for index, finding in enumerate(findings, start=1):
                title = self._finding_title(finding)
                lines.extend(
                    [
                        f"### {index}. {title}",
                        "",
                        f"- 风险等级: {finding.get('severity', '信息')}",
                        f"- 漏洞链接: {finding.get('url', '') or finding.get('target', '') or '待补充'}",
                        f"- 漏洞描述: {finding.get('description', '待补充')}",
                        f"- 测试过程: {finding.get('test_process') or finding.get('evidence') or '待补充'}",
                        f"- 风险分析: {finding.get('risk_analysis') or finding.get('impact') or '待补充'}",
                        f"- 修复建议: {finding.get('remediation', '待补充')}",
                    ]
                )
                screenshot_path = finding.get("screenshot_path", "")
                if screenshot_path:
                    lines.append(f"- 截图: {screenshot_path}")
                lines.append("")

        lines.extend(
            [
                "## 测试结论",
                "",
                "本次评估已基于 verified_findings artifact 生成正式报告。建议按风险等级优先修复高危和中危问题，并复测验证。",
                "",
            ]
        )
        return "\n".join(lines)

    def _build_html_report(
        self,
        target: str,
        findings: List[Dict],
        report_title: str,
        fallback_reason: str = "",
    ) -> str:
        start_time, end_time = self._report_times()
        severity_counts = self._count_severities(findings)
        target_profile_rows = self._target_profile_rows(target)
        target_rows_html = "".join(
            f"<tr><td>{escape(label)}</td><td>{escape(value)}</td></tr>"
            for label, value in target_profile_rows
        )
        summary_rows = "".join(
            f"<tr><td>{escape(level)}</td><td>{count}</td></tr>"
            for level, count in [
                ("严重", severity_counts["严重"]),
                ("高危", severity_counts["高危"]),
                ("中危", severity_counts["中危"]),
                ("低危", severity_counts["低危"]),
                ("信息", severity_counts["信息"]),
            ]
        )

        finding_blocks: list[str] = []
        for index, finding in enumerate(findings, start=1):
            screenshot_html = ""
            screenshot_path = finding.get("screenshot_path", "")
            if screenshot_path:
                screenshot_html = f"<p><strong>截图:</strong> {escape(str(screenshot_path))}</p>"
            finding_blocks.append(
                "\n".join(
                    [
                        "<section class=\"finding\">",
                        f"<h3>{index}. {escape(self._finding_title(finding))}</h3>",
                        f"<p><strong>风险等级:</strong> {escape(str(finding.get('severity', '信息')))}</p>",
                        f"<p><strong>漏洞链接:</strong> {escape(str(finding.get('url', '') or finding.get('target', '') or '待补充'))}</p>",
                        f"<p><strong>漏洞描述:</strong> {escape(str(finding.get('description', '待补充')))}</p>",
                        f"<p><strong>测试过程:</strong> {escape(str(finding.get('test_process') or finding.get('evidence') or '待补充'))}</p>",
                        f"<p><strong>风险分析:</strong> {escape(str(finding.get('risk_analysis') or finding.get('impact') or '待补充'))}</p>",
                        f"<p><strong>修复建议:</strong> {escape(str(finding.get('remediation', '待补充')))}</p>",
                        screenshot_html,
                        "</section>",
                    ]
                )
            )

        fallback_html = ""
        if fallback_reason:
            fallback_html = f"<p><strong>说明:</strong> {escape(fallback_reason)}</p>"

        findings_html = "\n".join(finding_blocks) if finding_blocks else "<p>当前未发现已验证漏洞。</p>"
        return "\n".join(
            [
                "<!DOCTYPE html>",
                "<html lang=\"zh-CN\">",
                "<head>",
                "<meta charset=\"utf-8\">",
                f"<title>{escape(report_title)}</title>",
                "<style>body{font-family:Arial,sans-serif;max-width:960px;margin:40px auto;padding:0 16px;line-height:1.6}table{border-collapse:collapse;width:100%}th,td{border:1px solid #ccc;padding:8px;text-align:left}.finding{margin:24px 0;padding:16px;border:1px solid #ddd;border-radius:8px}</style>",
                "</head>",
                "<body>",
                f"<h1>{escape(report_title)}</h1>",
                f"<p><strong>测试时间:</strong> {escape(start_time)} 至 {escape(end_time)}</p>",
                f"<p><strong>已验证发现数量:</strong> {len(findings)}</p>",
                fallback_html,
                "<h2>目标信息</h2>",
                "<table><thead><tr><th>字段</th><th>内容</th></tr></thead><tbody>",
                target_rows_html,
                "</tbody></table>",
                "<h2>风险汇总</h2>",
                "<table><thead><tr><th>风险等级</th><th>数量</th></tr></thead><tbody>",
                summary_rows,
                "</tbody></table>",
                "<h2>详细发现</h2>",
                findings_html,
                "<h2>测试结论</h2>",
                "<p>本次评估已基于 verified_findings artifact 生成正式报告。建议按风险等级优先修复高危和中危问题，并复测验证。</p>",
                "</body>",
                "</html>",
            ]
        )

    @tool()
    async def add_screenshot(
        self,
        name: Annotated[str, "截图名称，用于生成文件名"],
        full_page: Annotated[bool, "是否截取整页，默认是"] = True,
    ) -> str:
        """
        Capture a screenshot from the current browser page and save it into the workspace screenshots directory.
        """
        from toolset.browser import browser as browser_tool

        os.makedirs(self.screenshots_dir, exist_ok=True)
        context = await browser_tool.get_context()
        page = context.pages[0] if context.pages else await context.new_page()
        filename = f"{self._safe_filename(name)}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png"
        screenshot_path = os.path.join(self.screenshots_dir, filename)
        await page.screenshot(path=screenshot_path, full_page=full_page)
        return screenshot_path

    @tool()
    def generate_word_report(
        self,
        target: Annotated[str, "测试目标 URL"],
        findings: Annotated[List[Dict], "漏洞发现列表，每个漏洞包含 name, severity, description, evidence, remediation, screenshot_path"],
        report_title: Annotated[str, "报告标题"] = "Web应用渗透测试报告",
    ) -> str:
        """
        生成包含漏洞详情和截图的 Word 报告 (.docx)
        """
        try:
            from docx import Document
        except ImportError:
            return self.generate_markdown_report(
                target=target,
                findings=findings,
                report_title=report_title,
                fallback_reason="python-docx 不可用，已自动降级为 Markdown 报告。",
            )
        findings = self._normalize_findings(findings)
        doc = self._build_default_report(target, findings, report_title)

        filename = self._report_filename(target, "docx")
        report_path = os.path.join(self.workspace, filename)
        doc.save(report_path)
        self._write_final_report_reference(report_path, "docx", target, len(findings))
        return report_path

    @tool()
    def generate_markdown_report(
        self,
        target: Annotated[str, "测试目标 URL"],
        findings: Annotated[List[Dict], "漏洞发现列表；若传空列表则自动从 verified_findings artifact 读取"],
        report_title: Annotated[str, "报告标题"] = "Web应用渗透测试报告",
        fallback_reason: Annotated[str, "降级或补充说明"] = "",
    ) -> str:
        """
        生成 Markdown 报告 (.md)
        """
        findings = self._normalize_findings(findings)
        report_body = self._build_markdown_report(target, findings, report_title, fallback_reason)
        filename = self._report_filename(target, "md")
        report_path = os.path.join(self.workspace, filename)
        with open(report_path, "w", encoding="utf-8") as f:
            f.write(report_body)
        self._write_final_report_reference(report_path, "md", target, len(findings))
        return report_path

    @tool()
    def generate_html_report(
        self,
        target: Annotated[str, "测试目标 URL"],
        findings: Annotated[List[Dict], "漏洞发现列表；若传空列表则自动从 verified_findings artifact 读取"],
        report_title: Annotated[str, "报告标题"] = "Web应用渗透测试报告",
        fallback_reason: Annotated[str, "降级或补充说明"] = "",
    ) -> str:
        """
        生成 HTML 报告 (.html)
        """
        findings = self._normalize_findings(findings)
        report_body = self._build_html_report(target, findings, report_title, fallback_reason)
        filename = self._report_filename(target, "html")
        report_path = os.path.join(self.workspace, filename)
        with open(report_path, "w", encoding="utf-8") as f:
            f.write(report_body)
        self._write_final_report_reference(report_path, "html", target, len(findings))
        return report_path

    @tool()
    def generate_report(
        self,
        target: Annotated[str, "测试目标 URL"],
        findings: Annotated[List[Dict], "漏洞发现列表；若传空列表则自动从 verified_findings artifact 读取"],
        report_title: Annotated[str, "报告标题"] = "Web应用渗透测试报告",
        format: Annotated[str, "报告格式，支持 docx / md / html"] = "docx",
    ) -> str:
        """
        生成报告。优先支持 docx，必要时可使用 markdown 或 html。
        """
        normalized = format.lower()
        if normalized == "docx":
            return self.generate_word_report(target, findings, report_title)
        if normalized in {"md", "markdown"}:
            return self.generate_markdown_report(target, findings, report_title)
        if normalized == "html":
            return self.generate_html_report(target, findings, report_title)
        return "[ERROR] Unsupported report format. Use docx, md, or html."

    @tool()
    def generate_word_report_from_artifacts(
        self,
        target: Annotated[str, "测试目标 URL"],
        report_title: Annotated[str, "报告标题"] = "Web应用渗透测试报告",
    ) -> str:
        """
        从 verified_findings artifact 自动生成 Word 报告。
        """
        return self.generate_word_report(target=target, findings=[], report_title=report_title)

    @tool()
    def add_finding_with_screenshot(
        self,
        findings_list: Annotated[List, "漏洞列表（会被修改添加新漏洞）"],
        name: Annotated[str, "漏洞名称"],
        severity: Annotated[str, "严重程度: 严重/高危/中危/低危/信息"],
        description: Annotated[str, "漏洞描述"],
        evidence: Annotated[str, "证据/Payload"],
        remediation: Annotated[str, "修复建议"],
        screenshot_path: Annotated[Optional[str], "截图文件路径"] = None,
    ) -> List:
        """
        添加一个带截图的漏洞发现到列表
        """
        finding = {
            "name": name,
            "severity": severity,
            "description": description,
            "evidence": evidence,
            "impact": evidence,
            "remediation": remediation,
            "screenshot_path": screenshot_path,
            "type": self._guess_vuln_type(name),
        }
        findings_list.append(finding)
        return findings_list

    def _guess_vuln_type(self, name: str) -> str:
        """根据漏洞名称猜测类型"""
        name_lower = name.lower()
        type_keywords = {
            "sql": "SQL注入",
            "注入": "注入漏洞",
            "xss": "跨站脚本",
            "跨站": "跨站脚本",
            "上传": "文件上传",
            "下载": "任意文件下载",
            "遍历": "目录遍历",
            "包含": "文件包含",
            "rce": "远程代码执行",
            "命令": "命令执行",
            "越权": "权限绕过",
            "泄露": "信息泄露",
            "枚举": "信息泄露",
            "头部": "安全配置",
            "配置": "安全配置",
        }
        for keyword, vuln_type in type_keywords.items():
            if keyword in name_lower:
                return vuln_type
        return "其他"
